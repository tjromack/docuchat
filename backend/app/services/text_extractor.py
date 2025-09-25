import PyPDF2
from docx import Document as DocxDocument
from typing import Dict, Any
import io

class TextExtractor:
    """Extract text from various document formats"""
    
    @staticmethod
    def extract_text(file_content: bytes, file_type: str) -> Dict[str, Any]:
        """Extract text based on file type"""
        try:
            if file_type == '.pdf':
                return TextExtractor._extract_from_pdf(file_content)
            elif file_type == '.docx':
                return TextExtractor._extract_from_docx(file_content)
            elif file_type == '.txt':
                return TextExtractor._extract_from_txt(file_content)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
        except Exception as e:
            return {
                "success": False,
                "error": str(e),
                "text": "",
                "metadata": {}
            }
    
    @staticmethod
    def _extract_from_pdf(file_content: bytes) -> Dict[str, Any]:
        text = ""
        metadata = {"pages": 0}
        
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
        metadata["pages"] = len(pdf_reader.pages)
        
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        
        return {
            "success": True,
            "text": text.strip(),
            "metadata": metadata
        }
    
    @staticmethod
    def _extract_from_docx(file_content: bytes) -> Dict[str, Any]:
        doc = DocxDocument(io.BytesIO(file_content))
        text = ""
        
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        metadata = {
            "paragraphs": len(doc.paragraphs),
            "tables": len(doc.tables)
        }
        
        return {
            "success": True,
            "text": text.strip(),
            "metadata": metadata
        }
    
    @staticmethod
    def _extract_from_txt(file_content: bytes) -> Dict[str, Any]:
        try:
            text = file_content.decode('utf-8')
        except UnicodeDecodeError:
            text = file_content.decode('latin-1')
        
        metadata = {
            "lines": len(text.split('\n')),
            "characters": len(text)
        }
        
        return {
            "success": True,
            "text": text,
            "metadata": metadata
        }